from flask import (
    render_template,
    request,
    redirect,
    url_for,
    flash,
    send_from_directory,
    session,
    make_response,
    jsonify,
) 
import os 
import secrets
from io import BytesIO
from docx import Document
from datetime import datetime, timedelta

from ..config import Config
from ..models import PemeriksaanJob
from ..services.ekstraksi_teks_service import TextExtractor
from ..services.preprocessing_service import PreprocessingService
from ..services.pemeriksaan_dokumen_service import PemeriksaanDokumenService
from ..services.pemeriksaan_job_service import PemeriksaanJobService
from ..services.riwayat_service import RiwayatService
from ..services.auth_service import AuthService
from ..utils.file_utils import FileUtils
from ..utils.text_utils import TextNormalizer
from ..utils.docx_utils import DocxUtils


class SistemWeb:
    ERROR_TYPE_MAPPING = {
        "BD": "Tanda Baca Dasar",
        "DD": "Tanda Titik",
        "CmD": "Tanda Koma",
        "QtD": "Tanda Petik",
        "QsD": "Tanda Tanya",
        "HD": "Tanda Hubung",
        "CnD": "Titik Dua",
    }

    # Fungsi untuk menginisialisasi layanan pemeriksaan dokumen dengan opsi untuk menyuntikkan layanan preprocessing,
    # koreksi, dan aturan deteksi yang dapat disesuaikan, atau menggunakan default jika tidak diberikan.
    def __init__(
        self,
        pemeriksaan_service=None,
        auth_service=None,
        riwayat_service=None,
        preprocessing_service=None,
        docx_extractor=None,
        file_utils=None,
        docx_utils=None,
        text_normalizer=None,
        job_service=None,
    ):
        self.text_normalizer = text_normalizer or TextNormalizer()
        self.preprocessing_service = preprocessing_service or PreprocessingService(
            text_normalizer=self.text_normalizer
        )
        self.docx_extractor = docx_extractor or TextExtractor()
        self.file_utils = file_utils or FileUtils()
        self.docx_utils = docx_utils or DocxUtils()
        self.pemeriksaan_service = pemeriksaan_service or PemeriksaanDokumenService(
            preprocessing_service=self.preprocessing_service
        )
        self.auth_service = auth_service or AuthService()
        self.riwayat_service = riwayat_service or RiwayatService()
        self.job_service = job_service or PemeriksaanJobService(
            docx_extractor=self.docx_extractor,
            preprocessing_service=self.preprocessing_service,
            pemeriksaan_service=self.pemeriksaan_service,
            file_utils=self.file_utils,
            text_normalizer=self.text_normalizer,
        )

    # Fungsi untuk membersihkan file hasil pemeriksaan sebelumnya agar tidak menumpuk di server
    def _cleanup_current_result_files(self):
        current_file = session.get("current_file")
        if current_file:
            self.file_utils.remove_file_if_exists(Config.UPLOAD_FOLDER, current_file)
            self.file_utils.remove_file_if_exists(Config.UPLOAD_FOLDER, f"{current_file}.txt")
            self.file_utils.remove_file_if_exists(Config.UPLOAD_FOLDER, f"{current_file}.json")
            self.file_utils.remove_file_if_exists(Config.UPLOAD_FOLDER, f"{current_file}.sbd.json")
            self.file_utils.remove_file_if_exists(Config.UPLOAD_FOLDER, f"{current_file}.tokens.json")
            self.file_utils.remove_file_if_exists(Config.UPLOAD_FOLDER, f"{current_file}.pos.json")
            self.file_utils.remove_file_if_exists(
                Config.DETECTION_RESULT_FOLDER,
                f"{current_file}.txt",
            )
            self.file_utils.remove_file_if_exists(
                Config.DETECTION_RESULT_FOLDER,
                f"{current_file}.highlight.html",
            )
            self.file_utils.remove_file_if_exists(
                Config.CORRECTION_RESULT_FOLDER,
                f"{current_file}.txt",
            )
            self.file_utils.remove_file_if_exists(
                Config.CORRECTION_RESULT_FOLDER,
                f"{current_file}.correction.highlight.html",
            )
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, f"{current_file}.txt")
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, f"{current_file}.json")
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, f"{current_file}.sbd.json")
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, f"{current_file}.tokens.json")
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, f"{current_file}.pos.json")

        extracted_text_file = session.get("extracted_text_file")
        if extracted_text_file:
            self.file_utils.remove_file_if_exists(
                Config.DETECTION_RESULT_FOLDER,
                extracted_text_file,
            )

        detection_html_file = session.get("detection_result_html_file")
        if detection_html_file:
            self.file_utils.remove_file_if_exists(
                Config.DETECTION_RESULT_FOLDER,
                detection_html_file,
            )

        correction_result_file = session.get("correction_result_file")
        if correction_result_file:
            self.file_utils.remove_file_if_exists(
                Config.CORRECTION_RESULT_FOLDER,
                correction_result_file,
            )

        correction_result_html_file = session.get("correction_result_html_file")
        if correction_result_html_file:
            self.file_utils.remove_file_if_exists(
                Config.CORRECTION_RESULT_FOLDER,
                correction_result_html_file,
            )

        debug_normalized_file = session.get("debug_normalized_file")
        if debug_normalized_file:
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, debug_normalized_file)

        structured_text_file = session.get("structured_text_file")
        if structured_text_file:
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, structured_text_file)

        sbd_file = session.get("sbd_file")
        if sbd_file:
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, sbd_file)

        tokens_file = session.get("tokens_file")
        if tokens_file:
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, tokens_file)

        pos_file = session.get("pos_file")
        if pos_file:
            self.file_utils.remove_file_if_exists(Config.DEBUG_FOLDER, pos_file)

    def _clear_current_result_session(self):
        session["current_file"] = None
        session.pop("preview_filename", None)
        session.pop("extracted_text_file", None)
        session.pop("detection_result_html_file", None)
        session.pop("correction_result_file", None)
        session.pop("correction_result_html_file", None)
        session.pop("koreksi_text", None)
        session.pop("debug_normalized_file", None)
        session.pop("structured_text_file", None)
        session.pop("sbd_file", None)
        session.pop("tokens_file", None)
        session.pop("pos_file", None)
        session.pop("history_saved", None)
        session.pop("saved_history_id", None)
        session.pop("result_token", None)
        session.pop("current_job_id", None)
        session.pop("current_job_token", None)
        session["result_ready"] = False

    def _remember_job(self, job):
        jobs = session.get("jobs") or {}
        jobs[str(job.id)] = job.job_token
        session["jobs"] = jobs
        session["current_job_id"] = job.id
        session["current_job_token"] = job.job_token

    def _session_job_tokens(self):
        jobs = session.get("jobs") or {}
        if isinstance(jobs, dict):
            return {str(job_id): token for job_id, token in jobs.items()}

        # Compatibility for older session shapes.
        tokens = {}
        if session.get("current_job_id") and session.get("current_job_token"):
            tokens[str(session.get("current_job_id"))] = session.get("current_job_token")
        return tokens

    def _visible_jobs(self):
        now = datetime.utcnow()
        session_tokens = self._session_job_tokens()
        jobs_by_id = {}

        if session.get("user_id"):
            user_jobs = (
                PemeriksaanJob.query.filter(
                    PemeriksaanJob.pengguna_id == session.get("user_id"),
                    PemeriksaanJob.expires_at > now,
                )
                .order_by(PemeriksaanJob.created_at.desc(), PemeriksaanJob.id.desc())
                .all()
            )
            for job in user_jobs:
                jobs_by_id[job.id] = job

        session_job_ids = [int(job_id) for job_id in session_tokens.keys() if str(job_id).isdigit()]
        if session_job_ids:
            session_jobs = (
                PemeriksaanJob.query.filter(
                    PemeriksaanJob.id.in_(session_job_ids),
                    PemeriksaanJob.expires_at > now,
                )
                .order_by(PemeriksaanJob.created_at.desc(), PemeriksaanJob.id.desc())
                .all()
            )
            for job in session_jobs:
                if session_tokens.get(str(job.id)) == job.job_token:
                    jobs_by_id[job.id] = job

        return sorted(jobs_by_id.values(), key=lambda job: (job.created_at, job.id), reverse=True)

    def _has_visible_jobs(self):
        return bool(self._visible_jobs())

    def _clear_preview_and_results(self):
        self._cleanup_current_result_files()
        self._clear_current_result_session()

    def _build_document_preview(self, filename):
        if not filename:
            return None

        file_path = os.path.join(Config.UPLOAD_FOLDER, filename)
        if not os.path.exists(file_path):
            return {
                "paragraphs": [],
                "error": "Dokumen tidak ditemukan. Silakan unggah ulang.",
            }

        try:
            extract_result = self.docx_extractor.extract(file_path)
        except Exception:
            return {
                "paragraphs": [],
                "error": "Pratinjau dokumen belum bisa ditampilkan.",
            }

        paragraphs = extract_result.get("paragraphs") or []
        visible_paragraphs = [paragraph for paragraph in paragraphs if paragraph][:80]
        return {
            "paragraphs": visible_paragraphs,
            "is_truncated": len(paragraphs) > len(visible_paragraphs),
            "error": "",
        }

    def _can_save_current_result_to_history(self):
        return bool(
            session.get("user_id")
            and not session.get("history_saved")
            and session.get("result_token")
            and session.get("current_file")
            and session.get("detection_result_html_file")
            and session.get("correction_result_file")
            and session.get("correction_result_html_file")
        )

    def _save_current_result_to_history(self):
        if not self._can_save_current_result_to_history():
            return None

        riwayat = self.riwayat_service.simpan_dari_session(
            pengguna_id=session.get("user_id"),
            session_data=session,
            file_utils=self.file_utils,
        )
        if riwayat:
            session["history_saved"] = True
            session["saved_history_id"] = riwayat.id
        return riwayat

    # Endpoint untuk unggah dokumen dan tampilkan preview
    def unggah_dokumen(self):
        if request.method == "POST":
            if "file" not in request.files:
                flash("Tidak ada file yang diunggah.")
                return redirect(request.url)

            file = request.files["file"]
            if file.filename == "":
                flash("Nama file kosong.")
                return redirect(request.url)

            if file and self.docx_utils.allowed_file(file.filename, Config.ALLOWED_EXTENSIONS):
                old_file = session.get("current_file")
                if old_file and session.get("preview_filename") == old_file and not session.get("result_ready"):
                    self._cleanup_current_result_files()

                self._clear_current_result_session()

                safe_filename = self.docx_utils.secure_filename_safe(file.filename)
                filename = f"{secrets.token_hex(4)}_{safe_filename}"
                file_path = self.file_utils.remove_file_if_exists(
                    Config.UPLOAD_FOLDER,
                    filename,
                    return_path=True,
                )
                file.save(file_path)

                session["preview_filename"] = filename
                session["current_file"] = filename
                session["show_preview"] = True
                session["result_ready"] = False
                return redirect(url_for("main.upload_dokumen"))

            flash("File harus berformat DOCX.")

        filename = session.get("preview_filename")
        show_preview = session.pop("show_preview", False)
        preview_url = url_for("main.uploaded_file", filename=filename) if filename else None
        document_preview = self._build_document_preview(filename) if filename else None

        if not show_preview:
            self._save_current_result_to_history()

        response = make_response(
            render_template(
                "upload.html",
                preview_url=preview_url,
                document_preview=document_preview,
                filename=filename,
                display_filename=self._display_filename(filename),
                max_file_size_mb=Config.MAX_FILE_SIZE // (1024 * 1024),
                has_process_jobs=self._has_visible_jobs(),
            )
        )
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        return response
    
    # Fungsi tambahan untuk mengubah output POS tag dari Stanza menjadi format yang lebih sederhana untuk digunakan dalam pemeriksaan aturan
    def _flatten_pos_tags(self, pos_tags, normalized_text):
        """
        Konversi output POSTagger (list of list of dict) ke format flat
        yang dibutuhkan TitikRule: list of dict dengan key text, upos,
        start_char, end_char.
        """
        flat = []
        search_start = 0
        for sent in pos_tags:
            for tag in sent:
                word = tag["token"]
                # Cari posisi token di teks asli mulai dari search_start
                idx = normalized_text.find(word, search_start)
                if idx == -1:
                    # Kalau tidak ketemu, skip tapi jangan geser search_start
                    flat.append({
                        "text": word,
                        "upos": tag["upos"],
                        "xpos": tag["xpos"],
                        "lemma": tag.get("lemma", ""),
                        "start_char": -1,
                        "end_char": -1,
                    })
                    continue
                flat.append({
                    "text": word,
                    "upos": tag["upos"],
                    "xpos": tag["xpos"],
                    "lemma": tag.get("lemma", ""),
                    "start_char": idx,
                    "end_char": idx + len(word),
                })
                search_start = idx + len(word)
        return flat

    # Endpoint untuk menampilkan hasil deteksi dan koreksi
    def tampilkan_hasil(self):
        if request.method == "POST":
            current_file = session.get("current_file")
            if not current_file:
                flash("Unggah dokumen terlebih dahulu.")
                return redirect(url_for("main.upload_dokumen"))

            file_path = os.path.join(Config.UPLOAD_FOLDER, current_file)
            if not os.path.exists(file_path):
                flash("Dokumen tidak ditemukan, silakan unggah ulang.")
                return redirect(url_for("main.upload_dokumen"))

            job = self.job_service.buat_job(
                nama_dokumen=current_file,
                pengguna_id=session.get("user_id"),
            )
            self._remember_job(job)
            session.pop("preview_filename", None)
            session["current_file"] = None
            session["result_ready"] = False
            return redirect(url_for("main.job_status_page"))

        if not session.get("result_ready"):
            return redirect(url_for("main.upload_dokumen"))

        extracted_text_file = session.get("extracted_text_file")
        extracted_text = (
            self.file_utils.read_text_file(
                Config.DETECTION_RESULT_FOLDER,
                extracted_text_file,
            )
            if extracted_text_file
            else ""
        )
        detection_html_file = session.get("detection_result_html_file")
        detection_html = (
            self.file_utils.read_text_file(
                Config.DETECTION_RESULT_FOLDER,
                detection_html_file,
            )
            if detection_html_file
            else ""
        )
        correction_result_file = session.get("correction_result_file")
        correction_text = (
            self.file_utils.read_text_file(
                Config.CORRECTION_RESULT_FOLDER,
                correction_result_file,
            )
            if correction_result_file
            else ""
        )
        correction_result_html_file = session.get("correction_result_html_file")
        correction_html = (
            self.file_utils.read_text_file(
                Config.CORRECTION_RESULT_FOLDER,
                correction_result_html_file,
            )
            if correction_result_html_file
            else ""
        )

        if not any([extracted_text, detection_html, correction_text, correction_html]):
            flash("File hasil koreksi sudah tidak tersedia. Silakan proses dokumen ulang.")
            self._clear_current_result_session()
            return redirect(url_for("main.upload_dokumen"))
        
        # Baca error summary
        error_summary = {}
        current_file = session.get("current_file")
        if current_file:
            summary_filename = f"{current_file}.summary.json"
            summary_data = self.file_utils.read_json_file(
                Config.DETECTION_RESULT_FOLDER,
                summary_filename,
            )
            if summary_data:
                error_summary = self._normalize_error_summary(summary_data)
        
        response = make_response(
            render_template(
                "hasil.html",
                extracted_text=extracted_text,
                detection_html=detection_html,
                correction_text=correction_text,
                correction_html=correction_html,
                document_name=session.get("current_file"),
                auto_save_history=bool(session.get("user_id")),
                result_url=url_for("main.hasil_koreksi"),
                back_url=url_for("main.job_status_page"),
                back_label="Kembali",
                download_url=url_for("main.unduh_hasil_koreksi"),
                original_download_url=self._current_original_download_url(),
                before_unload_url="",
                error_summary=error_summary,
            )
        )

        return response

    def _normalize_error_summary(self, summary_data):
        if not isinstance(summary_data, dict):
            return {}

        total = int(summary_data.get("total") or 0)
        summary_items = summary_data.get("items") or []
        if summary_items:
            items = []
            by_type = {}
            for item in summary_items:
                if not isinstance(item, dict):
                    continue
                name = item.get("name") or "Lainnya"
                count = int(item.get("count") or 0)
                type_key = item.get("type_key") or self._summary_type_key_from_name(name)
                if count:
                    items.append({"name": name, "count": count, "type_key": type_key})
                    by_type[name] = by_type.get(name, 0) + count
            return {
                "total": total,
                "by_type": by_type,
                "by_code": summary_data.get("by_code") or {},
                "items": sorted(items, key=lambda item: item["name"]),
            }

        by_code = summary_data.get("by_code") or {}
        if by_code:
            by_type = {}
            for kode, info in by_code.items():
                count = info.get("count", 0) if isinstance(info, dict) else 0
                type_key = self._summary_type_key(kode)
                if type_key not in by_type:
                    by_type[type_key] = {
                        "name": self._summary_type_name(type_key),
                        "count": 0,
                    }
                by_type[type_key]["count"] += int(count or 0)
        else:
            by_type = {}
            for type_name, count in (summary_data.get("by_type") or {}).items():
                type_key = self._summary_type_key_from_name(type_name)
                by_type[type_key] = {
                    "name": type_name,
                    "count": int(count or 0),
                }

        items = [
            {"type_key": type_key, "name": info["name"], "count": info["count"]}
            for type_key, info in sorted(by_type.items(), key=lambda item: item[1]["name"])
            if info["count"]
        ]

        return {
            "total": total,
            "by_type": {info["name"]: info["count"] for info in by_type.values()},
            "by_code": by_code,
            "items": items,
        }

    @classmethod
    def _summary_type_key(cls, kode):
        kode = str(kode or "")
        for prefix in sorted(cls.ERROR_TYPE_MAPPING, key=len, reverse=True):
            if kode.startswith(prefix):
                return prefix
        return "unknown"

    @classmethod
    def _summary_type_name(cls, type_key):
        return cls.ERROR_TYPE_MAPPING.get(type_key, "Lainnya")

    @classmethod
    def _summary_type_key_from_name(cls, type_name):
        for type_key, name in cls.ERROR_TYPE_MAPPING.items():
            if name == type_name:
                return type_key
        return "unknown"

    def tampilkan_status_job(self, job_id=None):
        if job_id is not None and not self._get_accessible_job(job_id):
            flash("Proses pemeriksaan tidak ditemukan atau sesi sudah berakhir.")
            return redirect(url_for("main.upload_dokumen"))

        jobs = self._visible_jobs()
        return render_template(
            "job_status.html",
            jobs=[self._serialize_job(job) for job in jobs],
            status_url=url_for("main.jobs_status_json"),
            upload_url=url_for("main.upload_dokumen"),
        )

    def status_jobs_json(self):
        return jsonify({"jobs": [self._serialize_job(job) for job in self._visible_jobs()]})

    def status_job_json(self, job_id):
        job = self._get_accessible_job(job_id)
        if not job:
            return jsonify({"status": "not_found", "progress": 0}), 404

        return jsonify(
            {
                "id": job.id,
                "status": job.status,
                "progress": job.progress,
                "error_message": job.error_message,
                "result_url": url_for("main.hasil_job", job_id=job.id)
                if job.status == PemeriksaanJob.STATUS_DONE
                else None,
            }
        )

    def _serialize_job(self, job):
        return {
            "id": job.id,
            "nama_dokumen": job.nama_dokumen,
            "status": job.status,
            "progress": job.progress,
            "error_message": job.error_message,
            "can_cancel": job.status in {
                PemeriksaanJob.STATUS_PENDING,
                PemeriksaanJob.STATUS_PROCESSING,
            },
            "can_show_result": job.status == PemeriksaanJob.STATUS_DONE,
            "result_url": url_for("main.hasil_job", job_id=job.id),
            "cancel_url": url_for("main.batalkan_job", job_id=job.id),
            "created_at": self._format_wib(job.created_at),
        }

    @staticmethod
    def _format_wib(value):
        if not value:
            return ""
        return (value + timedelta(hours=7)).strftime("%d/%m/%Y %H:%M")

    def batalkan_job(self, job_id):
        job = self._get_accessible_job(job_id)
        if not job:
            return jsonify({"status": "not_found"}), 404

        if job.status in {
            PemeriksaanJob.STATUS_PENDING,
            PemeriksaanJob.STATUS_PROCESSING,
        }:
            job.status = PemeriksaanJob.STATUS_CANCELLED
            job.progress = 100
            job.error_message = "Pemeriksaan dibatalkan."
            from ..extensions import db

            db.session.commit()

        return jsonify({"status": job.status, "progress": job.progress})

    def tampilkan_hasil_job(self, job_id):
        job = self._get_accessible_job(job_id)
        if not job:
            flash("Proses pemeriksaan tidak ditemukan atau sesi sudah berakhir.")
            return redirect(url_for("main.upload_dokumen"))

        if job.status in {PemeriksaanJob.STATUS_FAILED, PemeriksaanJob.STATUS_CANCELLED}:
            flash(job.error_message or "Gagal memproses dokumen.")
            return redirect(url_for("main.upload_dokumen"))

        if job.status != PemeriksaanJob.STATUS_DONE:
            return redirect(url_for("main.job_status_page", job_id=job.id))

        self._load_job_result_to_session(job)
        return self.tampilkan_hasil()

    def _current_original_download_url(self):
        job_id = session.get("current_job_id")
        if not job_id:
            return None

        job = self._get_accessible_job(job_id)
        if not job:
            return None

        file_path = os.path.join(Config.UPLOAD_FOLDER, job.nama_dokumen)
        if not os.path.exists(file_path):
            return None

        return f"/jobs/{job.id}/dokumen-asli"

    def unduh_dokumen_asli(self, job_id):
        job = self._get_accessible_job(job_id)
        if not job:
            flash("Dokumen asli tidak ditemukan atau sesi sudah berakhir.")
            return redirect(url_for("main.upload_dokumen"))

        file_path = os.path.join(Config.UPLOAD_FOLDER, job.nama_dokumen)
        if not os.path.exists(file_path):
            flash("Dokumen asli sudah tidak tersedia.")
            return redirect(url_for("main.job_status_page"))

        return send_from_directory(
            Config.UPLOAD_FOLDER,
            job.nama_dokumen,
            as_attachment=True,
            download_name=self._display_filename(job.nama_dokumen),
        )

    @staticmethod
    def _display_filename(filename):
        parts = str(filename or "").split("_", 1)
        if len(parts) == 2 and len(parts[0]) == 8:
            return parts[1]
        return filename

    def _get_accessible_job(self, job_id):
        job = PemeriksaanJob.query.get(job_id)
        if not job:
            return None

        if session.get("user_id") and job.pengguna_id == session.get("user_id"):
            return job

        session_tokens = self._session_job_tokens()
        if session_tokens.get(str(job.id)) == job.job_token:
            return job

        return None

    def _load_job_result_to_session(self, job):
        self._remember_job(job)
        session["current_file"] = job.nama_dokumen
        session["preview_filename"] = job.nama_dokumen
        session["extracted_text_file"] = job.extracted_text_file
        session["detection_result_html_file"] = job.detection_result_html_file
        session["correction_result_file"] = job.correction_result_file
        session["correction_result_html_file"] = job.correction_result_html_file
        session["debug_normalized_file"] = job.debug_normalized_file
        session["structured_text_file"] = job.structured_text_file
        session["sbd_file"] = job.sbd_file
        session["tokens_file"] = job.tokens_file
        session["pos_file"] = job.pos_file
        session["history_saved"] = False
        session.pop("saved_history_id", None)
        session["result_token"] = job.result_token
        session["result_ready"] = True

    def login(self):
        return self.auth_service.login()

    def tampilkan_riwayat(self):
        pengguna_id = session.get("user_id")
        if not pengguna_id:
            return []
        return self.riwayat_service.ambil_riwayat_pengguna(pengguna_id)

    def simpan_hasil_ke_riwayat(self):
        if not session.get("user_id"):
            return {"status": "ignored", "saved": False}

        riwayat = self._save_current_result_to_history()
        self._clear_preview_and_results()
        return {"status": "ok", "saved": bool(riwayat)}


_sistem_web = SistemWeb()


def upload_dokumen():
    return _sistem_web.unggah_dokumen()


def hasil_koreksi():
    return _sistem_web.tampilkan_hasil()


def job_status_page(job_id=None):
    return _sistem_web.tampilkan_status_job(job_id)


def job_status_json(job_id):
    return _sistem_web.status_job_json(job_id)


def jobs_status_json():
    return _sistem_web.status_jobs_json()


def batalkan_job(job_id):
    return _sistem_web.batalkan_job(job_id)


def hasil_job(job_id):
    return _sistem_web.tampilkan_hasil_job(job_id)


def unduh_dokumen_asli(job_id):
    return _sistem_web.unduh_dokumen_asli(job_id)


def simpan_hasil_ke_riwayat():
    return _sistem_web.simpan_hasil_ke_riwayat()


def uploaded_file(filename):
    response = send_from_directory(Config.UPLOAD_FOLDER, filename)
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


def unduh_hasil_koreksi():
    correction_result_file = session.get("correction_result_file")
    if not correction_result_file:
        flash("Hasil koreksi tidak tersedia.")
        return redirect(url_for("main.upload_dokumen"))

    koreksi_text = _sistem_web.file_utils.read_text_file(
        Config.CORRECTION_RESULT_FOLDER,
        correction_result_file,
    )
    if not koreksi_text:
        flash("Hasil koreksi tidak tersedia.")
        return redirect(url_for("main.upload_dokumen"))

    # Generate DOCX in memory
    doc = Document()
    
    # Split text by lines and add to document
    lines = koreksi_text.split('\n')
    for line in lines:
        if line.strip():  # Only add non-empty lines
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()  # Add empty paragraph to preserve spacing
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create response with DOCX file
    response = make_response(buffer.getvalue())
    response.headers["Content-Disposition"] = "attachment; filename=hasil_koreksi.docx"
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


# Endpoint untuk membersihkan preview dan hasil terkait (opsional, bisa dipanggil via AJAX)
def clear_preview():
    _sistem_web._clear_preview_and_results()
    return {"status": "ok"}


def tentang_page():
    return_url = request.args.get("return_url") or url_for("main.hasil_koreksi")
    if not return_url.startswith("/") or return_url.startswith("//"):
        return_url = url_for("main.hasil_koreksi")
    return render_template(
        "tentang.html",
        from_page=request.args.get("from_page"),
        return_url=return_url,
    )
